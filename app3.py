# Full app setup in one script (modularized)
# Required Libraries
import streamlit as st
from PyPDF2 import PdfReader
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import RetrievalQA, LLMChain
from langchain.prompts import PromptTemplate
from langchain_google_genai import GoogleGenerativeAI
import os
import pandas as pd
import plotly.express as px
import uuid
import base64
import tempfile
import fitz  # PyMuPDF
from docx import Document
import google.generativeai as genai
from google.api_core.exceptions import InvalidArgument
from dotenv import load_dotenv

load_dotenv()

st.set_page_config(layout="wide")
st.title("📚 PDF QA App")

# Initialize session state for uploaded files
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []

# Initialize Gemini model
@st.cache_resource
def load_gemini_model():
    # You'll need to get an API key from Google AI Studio
    api_key = os.getenv("GOOGLE_API_KEY")
    
    if not api_key:
        api_key = st.text_input("Enter your Google API Key", type="password")
        if not api_key:
            st.warning("Please enter a Google API key to continue")
            st.stop()
    
    # Configure the Gemini model
    try:
        # Configure the genai module
        genai.configure(api_key=api_key)
        
        # Verify available models
        models = genai.list_models()
        available_models = [m.name for m in models]
        
        # Check which model is available and select the appropriate one
        gemini_model_name = None
        for model_option in ["gemini-1.5-pro", "gemini-pro", "gemini-1.0-pro"]:
            if any(model_option in model for model in available_models):
                gemini_model_name = model_option
                break
                
        if not gemini_model_name:
            st.error(f"No Gemini model found. Available models: {available_models}")
            st.stop()
            
        st.success(f"Using Gemini model: {gemini_model_name}")
        
        # Initialize the LangChain wrapper for Gemini
        llm = GoogleGenerativeAI(
            model=gemini_model_name,
            google_api_key=api_key,
            temperature=0.3,
            max_output_tokens=512
        )
        return llm
    except Exception as e:
        st.error(f"Error initializing Gemini model: {str(e)}")
        st.stop()

# Session state for chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "analytics" not in st.session_state:
    st.session_state.analytics = []

# File uploader
pdf_files = st.file_uploader("Upload one or more PDFs", type="pdf", accept_multiple_files=True)

# Store uploaded files in session state for later use
if pdf_files:
    st.session_state.uploaded_files = pdf_files

# Interactive PDF Viewer
with st.expander("📑 PDF Viewer", expanded=False):
    try:
        if st.session_state.uploaded_files:
            # Display the uploaded files in a selection box
            pdf_file_names = [uploaded_file.name for uploaded_file in st.session_state.uploaded_files]
            pdf_file_names.insert(0, "Select PDF File")
            selected_pdf = st.selectbox("Select a PDF to view", pdf_file_names)
            
            # Retrieve the selected PDF file
            selected_file = None
            for uploaded_file in st.session_state.uploaded_files:
                if uploaded_file.name == selected_pdf:
                    selected_file = uploaded_file
                    break
            
            # Display the selected PDF
            if selected_file and selected_pdf != "Select PDF File":
                st.subheader(f"Viewing PDF: {selected_pdf}")
                
                # Read PDF file
                selected_file.seek(0)  # Reset file pointer to start
                pdf_bytes = selected_file.read()
                selected_file.seek(0)  # Reset file pointer after reading
                
                # Encode the PDF file in base64 for displaying in iframe
                pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
                
                # Display the PDF file in an iframe using an HTML embed
                pdf_display = f'<iframe src="data:application/pdf;base64,{pdf_base64}" width="100%" height="600" type="application/pdf"></iframe>'
                st.markdown(pdf_display, unsafe_allow_html=True)
    except Exception as e:
        st.error(f"Error displaying PDF: {str(e)}")

question = st.text_input("Ask a question across PDFs")

# Helper: Save files to temp and chunk
def load_and_chunk(file):
    # Save file pointer position
    file_pos = file.tell()
    
    # Reset file pointer to start
    file.seek(0)
    
    try:
        reader = PdfReader(file)
        all_text, page_map = "", {}
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            page_map[i] = text
            all_text += f"\n[Page {i + 1}]\n{text}"
        
        splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_text(all_text)
        
        # Reset file pointer to original position
        file.seek(file_pos)
        
        return chunks, page_map
    except Exception as e:
        st.error(f"Error processing PDF {file.name}: {str(e)}")
        file.seek(file_pos)  # Reset file pointer even if there's an error
        return [], {}

# Helper: Create FAISS store
def embed_documents(chunks):
    # Use HuggingFace embeddings instead of OpenAI
    try:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
        return FAISS.from_texts(chunks, embeddings)
    except Exception as e:
        st.error(f"Error creating embeddings: {str(e)}")
        return None

# Helper: Display PDF Page (both methods available)
def show_pdf_page(file, page_num, use_iframe=False):
    # Save current position
    file_pos = file.tell()
    
    # Reset file pointer
    file.seek(0)
    
    try:
        if use_iframe:
            # Read the entire PDF
            pdf_bytes = file.read()
            # Encode the PDF file in base64 for displaying in iframe
            pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
            # Display the PDF file in an iframe with page number parameter
            pdf_display = f'<iframe src="data:application/pdf;base64,{pdf_base64}#page={page_num}" width="100%" height="500" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
        else:
            # Original method using PyMuPDF to render as image
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name
                
            # Open the saved PDF
            doc = fitz.open(tmp_path)
            
            # Validate page number
            if page_num < 1 or page_num > len(doc):
                st.error(f"Invalid page number: {page_num}. Document has {len(doc)} pages.")
                return
                
            page = doc.load_page(page_num - 1)
            pix = page.get_pixmap()
            img_path = tmp_path.replace(".pdf", f"_page{page_num}.png")
            pix.save(img_path)
            st.image(img_path, caption=f"Page {page_num}")
            
            # Clean up
            doc.close()
            try:
                os.unlink(img_path)
                os.unlink(tmp_path)
            except Exception as e:
                pass  # Silently handle cleanup errors
    except Exception as e:
        st.error(f"Error displaying PDF page: {str(e)}")
    finally:
        # Reset file pointer to original position
        file.seek(file_pos)

# Helper: Summarize
@st.cache_data
def summarize_doc(chunks, _llm):
    summary_prompt = PromptTemplate(
        input_variables=["context"],
        template="Summarize this document:\n{context}"
    )
    chain = LLMChain(llm=_llm, prompt=summary_prompt)
    
    # Join only a subset of chunks to avoid token limits
    full_text = " ".join(chunks[:5])  # Limiting to first 5 chunks
    
    try:
        return chain.run({"context": full_text})
    except Exception as e:
        st.error(f"Error during summarization: {str(e)}")
        return "Error: Document too large to summarize or API error. Try with fewer pages."


# Initialize model and DBs
try:
    llm = load_gemini_model()
    file_chunks, vector_dbs, page_maps = {}, {}, {}
    
    if pdf_files:
        with st.spinner("Processing PDF files..."):
            for file in pdf_files:
                chunks, page_map = load_and_chunk(file)
                if chunks:  # Only create db if chunks were successfully extracted
                    db = embed_documents(chunks)
                    if db:  # Only store if db was successfully created
                        file_chunks[file.name] = chunks
                        page_maps[file.name] = page_map
                        vector_dbs[file.name] = db
except Exception as e:
    st.error(f"Error loading model or processing files: {str(e)}")

# Document Summarization UI
if pdf_files and file_chunks:
    with st.expander("📄 Document Summarization"):
        summarize_option = st.selectbox("Select a document to summarize", 
                                      ["All"] + [f.name for f in pdf_files if f.name in file_chunks])
        if st.button("Summarize"):
            with st.spinner("Summarizing..."):
                try:
                    if summarize_option == "All":
                        for file in pdf_files:
                            if file.name in file_chunks:
                                summary = summarize_doc(file_chunks[file.name], llm)
                                st.subheader(file.name)
                                st.write(summary)
                    else:
                        f = next(f for f in pdf_files if f.name == summarize_option)
                        summary = summarize_doc(file_chunks[f.name], llm)
                        st.subheader(f.name)
                        st.write(summary)
                except Exception as e:
                    st.error(f"Error during summarization: {str(e)}")

# Question Answering UI
results = []
if question and vector_dbs:
    try:
        for fname, db in vector_dbs.items():
            qa = RetrievalQA.from_chain_type(llm=llm, retriever=db.as_retriever())
            
            try:
                result = qa({"query": question})
                answer = result['result']
                
                context_docs = db.similarity_search(question, k=1)
                if context_docs:
                    context = context_docs[0].page_content
                    
                    # Extract page number safely
                    page_num = "Unknown"
                    try:
                        page_num_match = context.split("[Page ")
                        if len(page_num_match) > 1:
                            page_num = page_num_match[1].split("]")[0]
                    except:
                        pass
                    
                    st.markdown(f"### 📘 {fname} (Page {page_num})")
                    #st.write(highlight_text(context, answer))
                    st.write(answer)
                    
                    
                    
                    st.session_state.chat_history.append({
                        "file": fname, 
                        "page": page_num, 
                        "question": question, 
                        "answer": answer
                    })
                    
                    st.session_state.analytics.append({
                        "file": fname,
                        "page": int(page_num) if page_num.isdigit() else 0,
                        "confidence": 0.9,
                        "question": question
                    })
                    
                    results.append((fname, page_num, question, answer))
            except Exception as e:
                st.error(f"Error processing question for {fname}: {str(e)}")
    except Exception as e:
        st.error(f"Error during question answering: {str(e)}")

# Chat History Panel
if st.session_state.chat_history:
    with st.expander("💬 Chat History"):
        for entry in st.session_state.chat_history[::-1]:
            st.markdown(f"**{entry['file']}** | Page {entry['page']}\n> {entry['question']}\n→ {entry['answer']}")

# Downloadable Report
if results:
    with st.expander("📥 Download Q&A Report"):
        docx = Document()
        docx.add_heading("PDF QA Report", 0)
        for fname, page, q, a in results:
            docx.add_paragraph(f"File: {fname} | Page: {page}", style="List Bullet")
            docx.add_paragraph(f"Q: {q}")
            docx.add_paragraph(f"A: {a}\n")
        
        try:
            docx_path = os.path.join(tempfile.gettempdir(), f"report_{uuid.uuid4()}.docx")
            docx.save(docx_path)
            with open(docx_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
                st.markdown(f"[Download DOCX Report](data:application/octet-stream;base64,{b64})", unsafe_allow_html=True)
            # Clean up temporary files
            try:
                os.unlink(docx_path)
            except:
                pass
        except Exception as e:
            st.error(f"Error creating downloadable report: {str(e)}")

# Analytics Dashboard
if st.session_state.analytics:
    with st.expander("📊 Analytics Dashboard"):
        df = pd.DataFrame(st.session_state.analytics)
        col1, col2 = st.columns(2)
        with col1:
            st.dataframe(df)
        with col2:
            try:
                fig = px.histogram(df, x="file", color="page", title="Answer Distribution by File")
                st.plotly_chart(fig, use_container_width=True)
            except Exception as e:
                st.error(f"Error generating analytics chart: {str(e)}")
        
        st.markdown("Use filters below to explore:")
        file_filter = st.selectbox("Filter by file", ["All"] + list(df["file"].unique()))
        if file_filter != "All":
            st.dataframe(df[df["file"] == file_filter])